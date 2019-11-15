# -*- coding: utf-8 -*-

import matplotlib.pyplot as plt
import pandas as pd
from sklearn.linear_model import LogisticRegression
import numpy as np
#import statsmodels.formula.api as sm
import warnings
from abc import ABCMeta, abstractmethod
import datetime
from dateutil.relativedelta import *
import gc
#import weakref
import copy
import calendar

try:
    from docx import Document
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_BREAK
except Exception:
    print('For report generation please install python-docx package.')
#from joblib import Parallel, delayed

# Created by Anna Goreva and Dmitry Yudochev


warnings.simplefilter('ignore')

plt.rc('font', family='Verdana')
plt.style.use('seaborn-darkgrid')
#get_ipython().magic(u'matplotlib inline')


gc.enable()

def doc_add_table(doc, records, headers = None, merge = None, color = None):
    '''
    Adds a table to an object of Document class

    Parameters
    ----------
    doc: a Document object to insert the table into
    records: matrix with table values
    headers: headers for columns in the table
    merge: list, coordinates of cells to merge, e.g. [(0, 1), (0, 3)] merges (0, 1), (0, 2) and (0, 3)
    color: dict of cells and colors, colors should be in RGB hex notation, {(row, column) : color}

    Returns
    ---------
    doc with table
    '''

    if merge is None:
        merge=[]
    if color is None:
        color={}

    table = doc.add_table(rows = 1, cols = len(records[0]))


    if headers is not None:
        if len(headers) != len(records[0]):
            print ('Headers do not match the columns!')
        else:
            for i in range(len(headers)):
                table.rows[0].cells[i].paragraphs[0].add_run(headers[i]).bold = True

    for s in records:
        row_cells = table.add_row().cells
        for i in range(len(s)):
            row_cells[i].text = s[i]

    if len(merge) == 2:
        if len(merge[0]) == 2 and len(merge[1]) == 2:
                table = merge_cells(table, merge)
        else:
            print ('doc_add_table error: merge size is unexpected')

    for coord in color:
        table.rows[coord[0]].cells[coord[1]]._tc.get_or_add_tcPr().append(parse_xml((r'<w:shd {} w:fill="' + color[coord] + '"/>').format(nsdecls('w'))))

    table.style = 'TableGrid'
    return doc



def merge_cells(table, merge):
    '''
    Merges cells in a table

    Parameters
    -----------
    table: a table to merge cells in
    merge: list of cells' coordinates

    Returns
    -----------
    table with merged cells
    '''
    a = table.cell(merge[0][0], merge[0][1])
    b = table.cell(merge[1][0], merge[1][1])
    a.merge(b)
    return table




def dates_to_interval(d):
    '''
    Returns the first day, the month and the year of the minimal date and the last date, the month and the year of the maximal date.

    Parameters
    ------------
    d: pd.Series/list of datetimes

    Returns
    -----------
    'startdate - enddate'
    '''
    min_date = min(d)
    max_date = max(d)
    return min_date.replace(day = 1).strftime('%d.%m.%Y') + ' - ' + max_date.replace(day = calendar.monthrange(max_date.year, max_date.month)[1]).strftime('%d.%m.%Y')




class SubReport(metaclass = ABCMeta):
    """
    Inerface for report parts
    """
    @abstractmethod
    def __init__(self):
        self.name = 'base name'


    @abstractmethod
    def build(self, report):
        '''
        Builds a part of report

        Parameters
        -------------
        report: Document() - a report to add text into

        Returns
        ---------
        report: report with additional information
        '''
        return report



class Title(SubReport):
    '''
    Class with title for the report
    '''
    def __init__(self, heading1 = None, heading2 = None, heading3 = None, pic = None, pic_width = 3):
        '''
        Initialization of the title

        Parameters
        -----------
        heading1: str; heading of the first level
        heading2: str, optional; heading of the second level
        heading3: str, optional; heading of the third level
        pic: str, optional; name of the file with corporate logo
        pic_width: int, width of the label

        '''
        self.name = 'title'
        self.heading1 = heading1
        self.heading2 = heading2
        self.heading3 = heading3
        if pic is None:
            print('No title picture specified. Skipping logo..')
        self.pic = pic
        self.pic_width = Inches(pic_width)
        if pic_width is None and pic is not None:
            print ('Error! Please set pic_width!')


    def build(self, report = None):
        if report is None:
            report = Document()

        if self.pic is not None:
            report.add_picture(self.pic, width = self.pic_width)

        if self.heading1 is not None:
            report.add_heading(self.heading1, 0)

        if self.heading2 is not None:
            heading2_paragraph = report.add_paragraph()
            heading2_paragraph.paragraph_format.space_before = Pt(64)
            heading2_run = heading2_paragraph.add_run(self.heading2)
            heading2_run.font.size = Pt(14)
            heading2_run.bold = True

        if self.heading3 is not None:
            heading3_run = report.add_paragraph().add_run(self.heading3)
            heading3_run.font.size = Pt(14)


        report.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

        return report





class Introduction(SubReport):
    '''
    Class with common introductory information for reports
    '''
    def __init__(self, structure = True, model_description = None, new_terms  = None):
        '''
        Parameters
        -----------
        structure: boolean, whether to include the paragraph with the document's structure
        model_description: str, user-defined model description for model annotation
        new_terms: str, new terms for the report
        '''
        if new_terms is None:
            new_terms=[]

        self.name = 'introduction'
        self.structure = structure
        self.model_description = model_description
        self.new_terms = new_terms
        if self.model_description is None:
            print ('Introduction speaks: Warning! No model description entered.')


    def build(self, report = None):
        if report is None:
            report = Document()

        # intro
        report.add_heading('Общие положения', level = 1)
        report.add_heading('Цель документа', level = 2)
        report.add_paragraph('Данный отчет создан с целью обеспечения непрерывности бизнес-процессов, сохранения и распространения внутренних экспертных знаний внутри Банка, а также для подтверждения обоснованности и корректности разработанной модели.' )
        report.add_paragraph('Для достижения этих целей Отчёт подробно описывает основные методологические подходы к моделированию, а также предпосылки и решения, которые принимались на каждом из этапов разработки модели.')

        # document structure
        if self.structure:
            report.add_heading('Структура документа', level = 2)
            report.add_paragraph('Данный отчёт имеет следующую структуру:')
            structure = ['Раздел «Структура модели» описывает общую логику получения и применения итогового скорингового балла для сделки',
                         'Раздел «Сбор данных» содержит описание процедуры формирования выборки для построения модели, а также описание подходов к фильтрации данных и обработке проблемных данных',
                         'В разделе «Формирование списка факторов и проведение однофакторного анализа» приводится логика отбора факторов для включения в модель на основе однофакторного анализа',
                         'В разделе «Многофакторный анализ» приводится логика отбора факторов для включения в модель на основе анализа корреляций и мультиколлинеарности',
                         'Раздел «Построение модели» посвящён расчёту весов факторов в итоговой модели']

            for s in structure:
                report.add_paragraph(s, style = 'List Bullet')


        # model annotation
        report.add_heading('Сводная информация по модели', level = 2)
        if self.model_description is not None:
            report.add_paragraph(self.model_description)

        report.add_paragraph('Сбор данных для модели:')
        report = doc_add_table(report, (('Запрос данных и источники данных' , 'Разработка перечня собираемых данных. Определение источников данных.'),
                                      ('Применение алгоритма формирования выборки', 'Очистка данных, общие принципы формирования выборки, агрегирование данных, разделение данных на выборку для построения модели и выборку для тестирования.')),
                            ('Этап', 'Описание'))
        report.add_paragraph('Сводное описание однофакторного анализа:')
        report = doc_add_table(report, (('Формирование списка факторов и этапы проведения однофакторного анализа', 'Определение списка факторов и их параметров, приведение непрерывных факторов к дискретной форме. \nРасчёт WOE для каждого значения / диапазона значений фактора по формуле \nWOE_i=ln ((N_G (i)+0.5)*N_B⁄N_G/ (N_B(i)+0.5)), \nгде N_G (i) и N_G – количество недефолтных клиентов в диапазоне и по всей выборке, соответственно,  а N_B (i) и N_B – количество дефолтных клиентов в диапазоне и по всей выборке, соответственно.'),
                                      ('Оценка дискриминирующей силы факторов', 'Оценка прогнозной силы каждого фактора путём расчёта показателей Gini и F-value.')),
                             ('Этап', 'Описание'))
        model_paragraph = report.add_paragraph('Сводное описание построения и калибровки модели:')
        model_paragraph.paragraph_format.space_before = Pt(8)
        report = doc_add_table(report,
                             (('Анализ корреляций', 'Анализ корреляций (парной взаимной зависимости между факторами) и анализ факторов на мультиколлинеарность (расчёт значений VIF – variance inflation factor).'),
                              ('Построение модели', 'Вычисление весов факторов путём построения логистической регрессии. '),
                              ('Тестирование модели на тестовой подвыборке', 'Модель была протестирована на тестовой подвыборке базового продукта, которая не использовалась в разработке модели.')),
                             ('Этап', 'Описание'))


        # terminology
        report.add_heading('Термины и сокращения', level = 1)
        terms = ['БКИ – Бюро Кредитных Историй',
                 'Выборка – набор сделок и их параметров, отвечающих заданным характеристикам и представляющим из себя часть анализируемого портфеля',
                 'Вероятность дефолта – вероятность наступления дефолта по сделке в течение одного года с даты присвоения/корректировки рейтинга',
                 'Дефолт – наличие по сделке просроченного платежа сроком более 90 дней по рассматриваемому продукту',
                 'Дискретные факторы – факторы с ограниченным количеством вариантов значений',
                 'Калибровка – 1. Соотнесение баллов и вероятности дефолта по сделке 2. Преобразование баллов с целью получения требуемого распределение скоринговых баллов по потоку',
                 'Многофакторный анализ – анализ комбинаций факторов с целью определения модели, наиболее точно дифференцирующей “плохие” и “хорошие” сделки',
                 'PD – величина вероятности дефолта',
                 'Модель оценки PD – модель оценки вероятности дефолта, рассчитывающая индивидуальный рейтинг сделки на основании отобранных факторов',
                 'Непрерывные факторы – факторы с неограниченным количеством вариантов значений',
                 'Однофакторный анализ – анализ данных по выборке сделок с целью определения предсказательной силы факторов и определения окончательного списка факторов',
                 'Преобразование факторов – замена значений факторов на расчётные величины (баллы, WOE), связанные с оценкой значения целевого признака, относящегося к значению фактора',
                 'Рейтинг – то же, что и кредитный рейтинг',
                 'Рейтинговая модель – модель, присваивающая рейтинг сделкам',
                 'Скоринговая модель – то же, что и модель присвоения рейтингов PD без возможности связывать значения показателей качества сделок (скоринговых баллов) со значениями вероятности дефолта',
                 'Скоринговый балл – значение показателя качества сделок с точки зрения вероятности их дефолта'] + self.new_terms

        for term in terms:
            report.add_paragraph(term, style = 'List Bullet')


        # model structure
        report.add_heading('Структура модели', level = 1)
        mstructure1 = ['Данная скоринговая модель представляет собой одномодульную линейную модель, то есть простую сумму баллов, определяемых в зависимости от значений используемых в модели факторов',
                      'Перечень данных состоит из списка факторов, которые потенциально могут иметь значимость для определения вероятности дефолта. Первоначально, на основании экспертного опыта и доступности данных был сформирован первичный список подобных факторов.',
                      'Факторы из данного списка были разбиты с помощью однофакторных деревьев решений на диапазоны однородного уровня рисков и затем преобразованы в Weight of Evidence (WOE). WOE для каждого значения / диапазона значений фактора определяется по формуле:',
                      ]
        for s in mstructure1:
            report.add_paragraph(s, style = 'List Bullet')
        try:
            report.add_picture('WOE.png', width=Inches(2))
        except Exception:
            print('WARNING! No WOE.png was found. Skipping WoE formula picture..')
        report.add_paragraph('где N_G (i) и N_G – количество недефолтных клиентов в диапазоне и по всей выборке, соответственно,  а N_B (i) и N_B – количество дефолтных клиентов в диапазоне и по всей выборке, соответственно.')

        mstructure2 = ['Рассчитаны показатели корреляции между факторами.',
                       'Многофакторный анализ заключался в выборе комбинации факторов и подборе соответствующих им весов, обеспечивающих наибольшую прогнозную силу и стабильность модели.',
                       'Разработанная модель применена к данным: проанализированы дискриминационные свойства модели на тестовой выборке.']

        for s in mstructure2:
            report.add_paragraph(s, style = 'List Bullet')


        return report



class DataCollection(SubReport):
    '''
    Contains information about data marts and samples.
    '''
    def __init__(self, DS, target_definition = 'Наличие просроченной задолженности длительностью более 90 дней в течение первых 12 месяцев жизни кредита', data_description = None,
                 target_shift = 12, entity_name = '', time_column = None, time_format = '%d.%m.%Y', list_features = True):
        '''
        Parameters
        ----------
        DS: DataSamples on which the model was built
        target_definition: str or None, by default target is considered a PD, if None then no target defition will be written, else - user-defined target definition
        target_shift: int, number of months after time in time_column when the targets are ready for analysis, e.g. for PD target_shift = 12, for income target_shift = 0
        entity_name: str, name of single sample in mart in Russian, in Genitive.
        time_column: str, name of column with time
        time_format: str, format of the time column
        list_features: to add or not to add the list of all features
        '''
        self.name = 'data collection'
        self.DS = DS
        self.target_definition = target_definition
        self.target_shift = target_shift
        self.entity_name = entity_name
        self.time_column = time_column
        self.time_format = time_format
        if self.entity_name is None or self.entity_name == '':
            print ('DataCollection speaking: no entity_name.')
        self.list_features = list_features
        self.data_description = data_description


    def add_date_intervals_table(self, col1 = '', add_months = 0):
        '''
        TECH

        Calculates date intervalds for train, validate and test

        Parameters
        ----------
        col1: str, text for the first column
        add_months: time shift in months

        Returns
        ----------
        A tow for pivot table
        '''
        if type(self.DS.train.dataframe[self.time_column].iloc[0]) == type(' '):
            train_dates = self.DS.train.dataframe[self.time_column].apply(lambda x: datetime.datetime.strptime(x, self.time_format))
            test_dates = self.DS.test.dataframe[self.time_column].apply(lambda x: datetime.datetime.strptime(x, self.time_format))
        else:
            train_dates = self.DS.train.dataframe[self.time_column]
            test_dates = self.DS.test.dataframe[self.time_column]
        t = ((col1, ) +
             (dates_to_interval(train_dates.apply(lambda x: x + relativedelta(months=add_months))), ) +
             (dates_to_interval(test_dates.apply(lambda x: x + relativedelta(months=add_months))), )
        )
        if self.DS.validate is not None:
            if type(self.DS.validate.dataframe[self.time_column].iloc[0]) == type(' '):
                val_dates = self.DS.validate.dataframe[self.time_column].apply(lambda x: datetime.datetime.strptime(x, self.time_format))
            else:
                val_dates = self.DS.validate.dataframe[self.time_column]
            t = t + (dates_to_interval(val_dates.apply(lambda x: x + relativedelta(months=add_months))), )
        return t



    def from_ds_stats(self, stats):
        '''
        TECH

        Adds to the table statistics from self.DS.stats()

        Parameters
        -----------
        stats: dictionary {name_of_statictics: index_in_stats}

        Returns
        ----------
        list of new rows
        '''
        result = []
        for (n, ind) in stats.items():
            t = (n, str(round(self.DS.stats()['Train'][ind], 4)), str(round(self.DS.stats()['Test'][ind], 4)), )
            if 'Validate' in self.DS.stats():
                t = t + (str(round(self.DS.stats()['Validate'][ind], 4)), )
            result.append(t)
        return result




    def add_pivot_table(self, report):
        '''
        Adds to report a pivot table with information abount data and samples

        Parameters
        ------------
        report: a Document to add the table to

        Returns
        -----------
        report: Document with table added
        '''
        headers = ('', 'Обучающая выборка', 'Выборка для тестирования')
        if self.DS.validate is not None:
            headers = headers + ('Выборка для валидации',)

        t_contents = []

        # time intervals
        if self.time_column is not None:
            t_contents.append(self.add_date_intervals_table('Временной период для формирования выборки ' + self.entity_name))
            t_contents.append(self.add_date_intervals_table('Горизонт выбора значений целевого признака', self.target_shift))

        # target definition
        if self.target_definition is not None:
            t = ('Определение целевого признака', self.target_definition, '', )
            if self.DS.validate is not None:
                t_contents.append(t + ('',))
            else:
                t_contents.append(t)
            if self.time_column is not None:
                merge = [(3, 1), (3, 3 - (self.DS.validate is None)*1)]
            else:
                merge = [(1, 1), (1, 3 - (self.DS.validate is None)*1)]
        else:
            merge = []

        # sample sizes
        t_contents = t_contents + self.from_ds_stats({'Количество ' + self.entity_name + ' в выборке':'amount',
                                                      'Количество целевых событий в выборке' : 'target',
                                                      'Количество нецелевых событий в выборке' : 'non_target',
                                                      'Доля целевых событий в выборке' : 'target_rate'})

        report = doc_add_table(report, t_contents, headers, merge = merge)


        return report



    def build(self, report = None):
        if report is None:
            report = Document()

        report.add_heading('Сбор данных', 1)
        report.add_heading('Сводная информация о выборках', 2)
        report = self.add_pivot_table(report)

        if self.list_features:
            report.add_heading('Исходные данные', 2)
            if self.data_description is not None:
                report.add_paragraph(self.data_description)
            report = doc_add_table(report, [(self.DS.train.target, '')] + [(i, '') for i in self.DS.train.features], ('Поле', 'Расшифровка'))
            print ('DataCollection speaks: please set the meanings of the features yourself.')

        report.add_heading('Преобразование данных', 2)
        report.add_paragraph('Исключены текстовые поля, а также id и даты заявок и договоров т.к. они непригодны для построения логистической регрессии, т.е. не являются факторами.', style = 'List Bullet')

        if self.DS.train.weights is not None:
            report.add_paragraph('Весами наблюдений служат значения ' + str(self.DS.train.weights) + '.')


        if len(self.DS.bootstrap) > 0:
            report.add_heading('Bootstrap-выборки', 2)
            report.add_paragraph('Сгенерировано ' + str(len(self.DS.bootstrap)) + ' bootstrap-выборок для проведения проверок стабильности факторов.')

        return report



class DataProcessing(SubReport):
    '''
    Class with information about missings processing and feature encodings
    '''
    def __init__(self, MP = None, FE = None):
        '''
        Parameters
        -----------
        DS: DataSamples() processed
        MP: MissingProcessor() object used
        FE: FEatureEncoder() object used
        '''
        self.name = 'data processing'
        self.MP = MP
        self.FE = FE


    def build(self, report = None):
        if report is None:
            report = Document()
        if self.MP is not None:
            mp_paragraph = report.add_paragraph('Обработка пропущенных значений', style = 'List Bullet')
            mp_paragraph.paragraph_format.space_before = Pt(8)
            report = self.write_missings(report)
        if self.FE is not None:
            fe_paragraph = report.add_paragraph('Кодирование категориальных факторов', style = 'List Bullet')
            fe_paragraph.paragraph_format.space_before = Pt(8)
            report = self.write_fe(report)
        return report


    def write_missings(self, report):
        '''
        TECH

        Writes down the table with missings processing

        Parameters
        -----------
        report: Document() to write into

        Returns
        ---------
        report with information about missings
        '''
        mp_stats = copy.deepcopy(self.MP.stats)
        mp_stats.action = mp_stats.action.apply(lambda x: str(x).replace('mean', 'Замена на среднее').replace('delete', 'Удаление из выборки').replace('distribution', 'Замена на значения из распределения'))

        return doc_add_table(report, [(mp_stats.T[f]['features'], mp_stats.T[f]['action'])for f in mp_stats.T],  ('Фактор', 'Способ обработки пропущенных значений'))


    def write_fe(self, report):
        '''
        TECH

        Writes down information abount feature encoding

        Parameters
        -----------

        '''
        fe_stats = copy.deepcopy(self.FE.stats)
        fe_stats.action = fe_stats.action.str.replace('{', 'Кодирование порядковыми номерами - ').replace('}', '').replace('one_hot', 'One-hot кодирование')
        return doc_add_table(report, [(fe_stats.T[f]['features'], fe_stats.T[f]['action'])for f in fe_stats.T],  ('Категориальный фактор', 'Способ кодирования'))




class OneFactor(SubReport):
    '''
    Class with information about WOE calculation and one-factor analysis
    '''
    def __init__(self, woe):
        '''
        Parameters
        -------------
        woe: a WOE object
        '''
        self.name = 'one factor analysis'
        self.stats = copy.deepcopy(woe.stats.groupby(by = 'feature').max().iteration.reset_index().merge(woe.stats, on = ['feature', 'iteration'])).drop(['Test'], 1)
        if woe.stats.shape[0] == 0:
            print ('OneFactor speaks: Warning: no information in woe.stats!')
        self.auto_fit_parameters = woe.auto_fit_parameters


    def build(self, report):
        if report is None:
            report = Document()

        report.add_heading('Однофакторный анализ и WOE-преобразование факторов', 1)
        report.add_paragraph('Для каждого фактора на обучающей выборке подобрано оптимальное разбиение значений на группы и произведён расчёт WOE для этих групп.')

        if self.auto_fit_parameters != {}:
            # feature selection was performed -> needs to be presented in the report
            report.add_paragraph('Для каждого фактора после подбора оптимальных групп значений для расчёта WOE были проведены следующие проверки, по которым принималось решение об исключении фактора из дальнейшего анализа:')

            # gini
            if self.auto_fit_parameters['G_on']:
                p = report.add_paragraph(style = 'List Bullet')
                p.add_run('Проверка значений Gini. ').bold = True
                p.add_run('Значения Gini на обучающей' + ('Validate' in self.stats)*', вылидационной' + ' и bootstrap выборках должны превышать '
                          + str(self.auto_fit_parameters['G_gini_threshold']) + '. Кроме того, разница значений Gini на обучающей'
                                     + ('Validate' in self.stats)*(', валидационной') + ' и bootstrap выборках не должна превышать '
                                     + str(self.auto_fit_parameters['G_gini_decrease_threshold'])
                                     + ('Validate' in self.stats)*(self.auto_fit_parameters['G_gini_increase_restrict'])*(', и Gini на валидационной выборке не должно превышать Gini на обучающей выборке более, чем на '
                                       + str(self.auto_fit_parameters['G_gini_decrease_threshold'])
                                     + '.'))

            # business logic
            if self.auto_fit_parameters['BL_on']:
                p = report.add_paragraph(style = 'List Bullet')
                p.add_run('Проверка бизнес-логики. ').bold = True
                to_write = 'По умолчанию для каждого интервального фактора проверяется монотонность WOE-преобразования. '
                if self.auto_fit_parameters['BL_conditions_dict'] is not None:
                    to_write = to_write + 'Дополнительные условия - в приложении.'
                    print ('Please attach to the report the file with BL_conditions_dict')
                p.add_run(to_write)

            # restrictions of WOE-groups
            if self.auto_fit_parameters['SM_on'] or self.auto_fit_parameters['WOEM_on']:
                p = report.add_paragraph(style = 'List Bullet')
                p.add_run('Ограничения при подборе разбиений на группы. ').bold = True
                p.add_run(self.auto_fit_parameters['SM_on']*('Минимальное количество наблюдений в группе: ' + str(self.auto_fit_parameters['SM_size_threshold'])
                    + ', минимальное количество целевых событий в группе: ' + str(self.auto_fit_parameters['SM_target_threshold']) + '. ')
                    + self.auto_fit_parameters['WOEM_on']*('Минимальная разница по WOE между группами: ' + str(self.auto_fit_parameters['WOEM_woe_threshold'])
                    + ', при проверках' + (self.auto_fit_parameters['WOEM_with_missing'] == False)*' не' + ' использовались в том числе и значения WOE для пропущенных значений.')
                    )

            # WOE order checks
            if self.auto_fit_parameters['WOEO_on']:
                p = report.add_paragraph(style = 'List Bullet')
                p.add_run('Проверка стабильности значений целевого признака по группам на bootstrap-выборках. ').bold = True
                p.add_run('Тренд должен быть стабилен для ' + str(self.auto_fit_parameters['WOEO_correct_threshold']*100) + '% bootstrap-выборок.')
        else:
            print('WARNING! Auto fit parameters are missing! This may occur after importing groups. You should fill this info manually or rerun autofit. Skipping..')

        report.add_paragraph('По каждому фактору были подобраны оптимальные группы для расчёта WOE, рассчитаны WOE-значения и коэффициенты Gini, где применимо.')

        self.stats['result'] = self.stats.result.apply(lambda x: 'Не исключается' if 'Success' in str(x) else 'Исключается')
        self.stats['reason'] = self.stats.reason.apply(lambda x: 'Не удалось разбить на группы по WOE' if 'After the attempt' in str(x) else
                  ('Фактор не прошёл проверку по Gini' if 'Gini check failed' in str(x) else
                   ('Фактор не прошёл проверку на соответствие бизнес-логике' if 'Business logic check failed' in str(x) else
                    ('Фактор не прошёл проверку на порядок WOE в bootstrap-выборках' if 'WoE order check' in str(x) else ''))))

        headers = ('Фактор', 'Результат отбора', 'Причина исключения', 'Gini на обучающей выборке', )
        if 'Validate' in self.stats:
            headers = headers + ('Gini на валидационной выборке', )
            self.stats.Validate = self.stats.Validate.apply(lambda x: 'нет' if pd.isnull(x) else str(round(x, 2)))
        self.stats.Train = self.stats.Train.apply(lambda x: 'нет' if pd.isnull(x) else str(round(x, 2)))


        records = []
        for ind, row in self.stats.drop(['iteration'], 1).iterrows():
            #records.append(list(row))
            if 'Validate' in self.stats:
                records.append([row.feature, row.result, row.reason, row.Train, row.Validate])
            else:
                records.append([row.feature, row.result, row.reason, row.Train])

        report = doc_add_table(report, records = records, headers = headers)

        # Manual changes in WOE groups
        if 'Train_final' in self.stats:
            paragraph_add = report.add_paragraph('По следующим факторам была проведена дополнительная корректировка разбиения на группы для расчёта WOE-значений:')
            paragraph_add.paragraph_format.space_before = Pt(8)

            stats_add = copy.deepcopy(self.stats[self.stats.Train_final.isnull() == False].drop(['iteration'], 1))
            headers_add = ('Фактор', 'Gini на обучающей выборке', )
            if 'Validate_final' in self.stats:
                headers_add = headers_add + ('Gini на валидационной выборке', )
                stats_add.Validate_final = stats_add.Validate_final.apply(lambda x: 'нет' if pd.isnull(x) else str(round(x, 2)))

            stats_add.Train_final = stats_add.Train_final.apply(lambda x: 'нет' if pd.isnull(x) else str(round(x, 2)))

            records_add = []
            for ind, row in stats_add.iterrows():
                if 'Validate_final' in self.stats:
                    records_add.append([row.feature, row.Train_final, row.Validate_final])
                else:
                    records_add.append([row.feature, row.Train_final])

            report = doc_add_table(report, records = records_add, headers = headers_add)

        return report




class MultiFactor(SubReport):
    '''
    Class with information about correlations and variance inflation factor
    '''
    def __init__(self, CA, vif = None, pic_width = 6):
        '''
        Parameters
        -----------
        CA: a CorrelationAnalyzer object
        vif: a VIF object
        pic_width: picture width in inches
        '''
        self.name = 'multi-factor analysis'
        self.CA = CA
        self.vif = vif
        self.pic_width= pic_width



    def build(self, report):
        if report is None:
            report = Document()

        corr_width = Inches(self.pic_width)
        report.add_heading('Многофакторный анализ', 1)
        report.add_heading('Анализ корреляций', 2)
        report.add_paragraph('Анализ корреляций между факторами охватывает анализ парной взаимной линейной зависимости между факторами и анализ наличия мультиколлинеарности. Целью анализа корреляций является исключение факторов с высокой корреляцией, поскольку использование коррелированных факторов в регрессии повышает стандартные отклонения оценок весов в многофакторном анализе, что снижает устойчивость и надёжность моделей.', style = 'List Bullet')

        method = 'Спирмена' if list(self.CA.stats.method)[0] == 'spearman' else 'Пирсона'

        report.add_paragraph('Для анализа взаимозависимости были рассчитаны корреляционные матрицы ' + method + '. Порог, начиная с которого коэффициент корреляции считается высоким, зафиксирован на уровне ' + str(list(self.CA.stats.threshold)[0]) + '. ', style = 'List Bullet')
        if self.CA.stats.out_before.isnull().sum() == 0:
            report.add_paragraph('Матрица корреляций факторов: ')
            report.add_picture(list(self.CA.stats['out_before'])[0], width=corr_width)

        report.add_paragraph('Матрица корреляций после удаления факторов с высокой корреляцией:', style = 'List Bullet')
        report.add_picture(list(self.CA.stats['out_after'])[0], width=corr_width)

        if self.vif is not None:
            report.add_paragraph('Таблица значений VIF факторов:', style = 'List Bullet')
            report.add_picture(list(self.vif.stats['out'])[0], width=corr_width)

        return report






class Stability(SubReport):
    '''
    Class with PSI reports.
    '''
    def __init__(self, SA, excluded_features = None):
        '''
        Parameters
        ----------
        SA: StabilityAnalyzer object
        excluded_features: list of features that were excluded by user due to lack of stability
        '''

        if excluded_features is None:
            excluded_features=[]

        self.name = 'stability analysis'
        self.SA = SA
        self.excluded_features = excluded_features if isinstance(excluded_features, list) else list(excluded_features)



    def build(self, report):
        if report is None:
            report = Document()

        report.add_heading('Анализ стабильности', 1)
        report.add_paragraph('Дополнительно был проведён анализ стабильности факторов, отобранных на этапе построения пробной модели. В качестве базового периода, с которым производилось сравнение, был выбран ноябрь 2015 года (обучающая выборка). Анализ стабильности проводился путём расчёта индексов стабильности.')
        try:
            report.add_picture('PSI.png', width=Inches(2))
        except Exception:
            print('WARNING! No PSI.png was found. Skipping PSI formula picture..')
        report.add_paragraph('где PSI (population stability index) – индекс стабильности, Ф%, О% - доля наблюдений соответствующей группы в анализируемом и базовом периодах, соответственно.')
        psi_headers = ('Диапазон значений', 'Расшифровка')
        psi_records = (('PSI < 0.1', 'Популяция стабильна, распределения не различаются'),
                       ('0.1 ≤ PSI < 0.25', 'Есть небольшие изменения в популяции, есть сомнения в том, что распределения не различаются'),
                       ('PSI ≥ 0.25', 'Значительные изменения в популяции, распределения различны'))
        report = doc_add_table(report, psi_records, psi_headers, color = {(1, 0) : '008F00', (2, 0) : 'FFFF00', (3, 0) : 'FF0000'})

        report.add_paragraph('Анализ стабильности отдельных факторов проводился в разрезе групп, полученных в результате однофакторного анализа.')

        for sample_name in self.SA.stats.sample_name.drop_duplicates():
            to_write = 'Стабильность на ' + ('обучающей' if 'train' in sample_name.lower() else ('валидационной' if 'valid' in sample_name.lower() else ('тестовой' if 'test' in sample_name.lower() else ''))) + ' выборке:'
            report.add_paragraph(to_write)
            sample_stats = self.SA.stats.loc[self.SA.stats.sample_name == sample_name]
            if list(sample_stats[sample_stats.parameter == 'out'].meaning)[0] == 1 and not pd.isnull(list(sample_stats[sample_stats.parameter == 'out_images'].meaning)[0]):
                report.add_picture(list(sample_stats[sample_stats.parameter == 'out_images'].meaning)[0] + 'stability.png', width = Inches(7))


        if len(self.excluded_features) > 0:
            report.add_paragraph('В связи с нестабильностью были удалены факторы ' + [f for f in self.excluded_features] + '.')

        return report



class ModelReport(SubReport):
    '''
    Class with information about scoring model
    '''
    def __init__(self, model, DS_woe, data_calibration=None, coefs_pic = 'model_coefficients.png', round_digits = 4):
        '''
        Parameters
        -----------
        model: a ScoringModel object
        DS_woe: Data or DataSamples to calculate model quality on
        data_calibration: Data or DataSamples object to calibrate model on. If None, then DS_woe.train will be used
        coefs_pic: name of file with model coefficients picture
        round_digits: rounding parameter
        '''
        self.name = 'model'
        self.model = model
        self.coefs_pic = coefs_pic
        self.DS_woe = DS_woe
        if data_calibration is None:
            print('Attention! No DataSamples\Data object for calibration specified. Using DS_woe.train instead..')
            self.data_calibration = DS_woe.train
        else:    
            if type(data_calibration) == Datasamples: 
                self.data_calibration = data_calibration.train
            elif type(data_raw) == Data:
                self.data_calibration = data_calibration
        self.round_digits = round_digits


    def scorecard_to_tab(self, report):
        headers = ('Фактор', 'Интервал значений', 'Пропущенные значения', 'Балл', 'Доля выборки', 'Доля целевых событий')
        scorecard = copy.deepcopy(self.model.scorecard)
        trees=scorecard[scorecard.apply(lambda row: pd.isnull(row['categorical']) and row['feature']!='intercept', axis=1)]['feature'].unique().tolist()
        scorecard.loc[scorecard.feature.isin(trees), 'values'] = scorecard.loc[scorecard.feature.isin(trees), 'values'].apply(lambda x:
                                                                    eval(x.replace(': nan,',': np.nan,').replace(': nan}',': np.nan}')\
                                                                          .replace('), nan)','), np.nan)').replace(', nan,',', np.nan,')\
                                                                          .replace('[nan,','[np.nan,').replace(', nan]',', np.nan]').replace('[nan]','[np.nan]')\
                                                                          .replace('(inf,','(np.inf,').replace(', inf)',', np.inf)')\
                                                                          .replace('(-inf,','(-np.inf,').replace(', -inf)',', -np.inf)')))


        scorecard.loc[scorecard.feature.isin(trees)==False, 'values']= \
            scorecard.loc[scorecard.feature.isin(trees)==False, 'values'].apply(lambda x: 'Пропущенное значение' if str(x)=='nan' else x)
        scorecard.loc[scorecard.feature.isin(trees)==False, 'values'] = \
            scorecard.loc[scorecard.feature.isin(trees)==False, 'values'].apply(lambda x: str(x).replace('[-inf,', '<').replace(', inf]', ''))
        scorecard.loc[scorecard.feature.isin(trees)==False, 'values'] = \
            scorecard.loc[scorecard.feature.isin(trees)==False, 'values'].apply(lambda x: x[:-1] if '<' in x else ('≥ ' + x[1:] if '[' in x and ']' not in x else(x)))
        scorecard.loc[scorecard.categorical == False, 'values'] = scorecard.loc[scorecard.categorical == False, 'values'].apply(lambda x: x.replace(']', ')'))

        # code for cross-features
        trees_conditions={}
        for ind in scorecard.loc[scorecard.feature.isin(trees), 'values'].index:
            conditions=scorecard.loc[scorecard.feature.isin(trees), 'values'][ind]
            str_conditions=''
            for ci in range(len(conditions)):
                for var in conditions[ci]:
                    if isinstance(conditions[ci][var], list):
                        if [np.nan] == conditions[ci][var]:
                            str_conditions+='(MISSING(' + var + '))'
                        else:
                            if np.nan in conditions[ci][var]:
                                str_conditions+='(MISSING(' + var + ') or '
                            else:
                                str_conditions+='(NOT MISSING(' + var + ') and '
                            str_conditions+=var + ' in ' + str([x for x in conditions[ci][var] if pd.isnull(x)==False])\
                                                                      .replace("[", "(").replace("]", ")") + ')'
                    else:
                        if pd.isnull(conditions[ci][var]):
                            str_conditions+='(MISSING(' + var + '))'
                        elif conditions[ci][var]==(-np.inf, np.inf):
                            str_conditions+='(NOT MISSING(' + var + '))'
                        else:
                            if pd.isnull(conditions[ci][var][1]):
                                str_conditions+='(MISSING(' + var + ') or '
                            else:
                                str_conditions+='(NOT MISSING(' + var + ') and '
                            if isinstance(conditions[ci][var], tuple):
                                cleared = conditions[ci][var][0] if pd.isnull(conditions[ci][var][1]) else conditions[ci][var]

                                str_conditions+=((str(cleared[0]) + ' <= ') if cleared[0]!=-np.inf else '')\
                                                 + var\
                                                 + (' < ' + str(cleared[1]) if cleared[1]!=np.inf else '') + ')'
                    str_conditions+=' and ' if list(conditions[ci]).index(var)!=len(conditions[ci])-1 else ''

                str_conditions+=' or ' if ci!=len(conditions)-1 else ''
            trees_conditions[ind]=str_conditions
        scorecard.loc[scorecard.feature.isin(trees), 'values']=pd.Series(trees_conditions)


        records = [('Начальный балл', '', '', str(list(scorecard[scorecard.feature == 'intercept'].score)[0]), '', '')]
        for (ind, row) in scorecard.iterrows():
            if ind > 0:
                records.append((row.feature, row['values'], str(row.missing), str(row.score), str(round(row.sample_part, self.round_digits)), str(round(row.ER, self.round_digits))))

        return doc_add_table(report, records, headers)



    def build(self, report):
        if report is None:
            report = Document()

        report.add_heading('Построение скоринговой модели', 1)
        if isinstance(self.model.model, LogisticRegression):
            report.add_paragraph('Построена логистическая регрессия c регуляризацией по норме ' + self.model.model.penalty +
                                 ' и коэффициентом регуляризации ' + str(round(1/self.model.model.C, self.round_digits)) + '.')
        try:
            if len(self.model.selected) > 0:
                report.add_paragraph('После проведения проверки факторов на значимость и соответствие бизнес-логике (отрицательные коэффициенты при WOE) были отобраны факторы: ' + str(self.model.selected)[1: -1].replace("'", "") + '.')
        except Exception:
            pass

        try:
            report.add_paragraph('Коэффициенты модели:')
            if self.coefs_pic is not None:
                report.add_picture(self.coefs_pic, width = Inches(5))

            headers = ('Фактор', 'Коэффициент')
            records = []
            for k, v in self.model.coefs.items():
                records.append((k, str(round(v,3))))
            report = doc_add_table(report, records, headers)
        except Exception:
            print ('ModelReport speaking: no model coefficients available for writing.')

        report.add_heading('Качество скоринговой модели', 2)

        gini = self.model.roc_curve(self.DS_woe, filename = 'roc_curve_report', verbose = False)
        with_validate = len(gini) == 3
        gini_paragraph = report.add_paragraph()
        gini_paragraph.add_run('Коэффициент Gini на обучающей выборке составляет ')
        gini_paragraph.add_run('Gini = ' + str(gini[0])).bold = True
        if with_validate:
            gini_paragraph.add_run(', качество на валидационной выборке составляет ')
            gini_paragraph.add_run('Gini = ' + str(gini[1])).bold = True
        gini_paragraph.add_run(', качество на тестовой выборке составляет ')
        gini_paragraph.add_run('Gini = ' + str(gini[-1]) + '.').bold = True

        report.add_paragraph('ROC-кривые:')
        report.add_picture('roc_curve_report.png', width = Inches(5))

        report.add_heading('Скоринговая карта', 2)
        if self.model.scorecard.shape[0] > 0:
            report = self.scorecard_to_tab(report)

        try:
            df_scored = self.model.score(self.data_calibration, features_to_leave=[self.data_calibration.target]).dataframe
            logreg = LogisticRegression(random_state = 42, solver='newton-cg', C=1000000)
            logreg.fit(df_scored[['score']], df_scored[self.data_calibration.target])

            report.add_heading('Коэффициент логистической регрессии', 2)
            report.add_paragraph('Для определения вероятности целевого события на основе скорингового балла, рассчитанного данной моделью, используются следующие параметры логистической регрессии: свободный член '
                                 + str(round(logreg.intercept_[0], self.round_digits)) + ' и коэффициент при скоринговом балле '
                                 + str(round(logreg.coef_[0][0], self.round_digits)) + '.')

        except Exception:
           print('WARNING! No initial data for calibration was found. Please, specify a DataSamples object with non-transformed data in DS_raw parameter. Skipping calibration..')


        return report




class Report:
    '''
    Class for generating a model construction report.
    '''

    def __init__(self):
        # Добавить стили, шрифты
        self.report = Document()


    def make_report(self, report_name, subreports = None):
        '''
        Makes a report in .docx format according to all settings.

        Parameters
        -----------
        report_name: name of file with the report
        subreports: list or set of objects SubReport() that contain parts of the current report

        '''

        self.report = Document()

        if subreports is None:
            subreports=[]

        if len(subreports) > 0:
            for sub in subreports:
                print ('Writing ' + sub.name + '...')
                self.report = sub.build(self.report)

        self.report.save(report_name + '.docx')
        print ('Your report is ready. Enjoy!')

